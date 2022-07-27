#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import os
import re
import docx
import PyPDF2
import io 
import nltk
import altair as alt
import numpy as np
nltk.download('stopwords')
nltk.download('punkt')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize 
import pandas as pd
import streamlit as st
import numpy as np
# loading the trained model
#pickle_in = open('classifier.pkl', 'rb') 
#classifier = pickle.load(pickle_in)


# Define key terms dictionary for fixing Role Applied for 
terms = {'WorkDay ERP':['workday', 'workday consultant', 'workday hcm', 'eib', 'picof', 
                        'workday studio','nnbound/outbound integrations'],
         'Peoplesoft':['peoplesoft', 'pia','ccb','birt','peci','ccw','pum','people tools',
                        'peoplesoft implementation','peoplesoft components',
                        'peoplesoft dba','peoplesoft admin','peoplesoft admin/dba','peopleSoft fscm', 
                        'peopletoolsupgrade','peopletools upgrade','process scheduler servers',
                        'peoplesoft hrms','peopleSoft consultant','peopledoft cloud',
                        'PeopleSoft migrations','eoplesoft Testing Framework','pure internet architecture'],             
         'Database Developer':['sql','sql server', 'ms sql server','msbi', 'sql developer', 'ssis','ssrs',
                        'ssms','t-sql','tsql','Razorsql', 'razor sql','triggers','powerbi','power bi',
                        'oracle sql', 'pl/sql', 'pl\sql','oracle', 'oracle 11g','oledb','cte','ddl',
                        'dml','etl','mariadb','maria db'],
         'Java Developer':['reactjs', 'react js', 'react js developer', 'html', 
                        'css3','xml','javascript','html5','boostrap','jquery', 'redux','php', 'node js',
                        'nodejs','apache','netbeans','nestjs','nest js','react developer','react hooks',
                        'jenkins']}

# List of all key terms to indicate skillset. Include all the key words in the list 
allTerms = ['workday', 'hcm', 'eib', 'picof','workday hcm',
                        'workday studio','nnbound/outbound integrations',
                        'peoplesoft', 'pia','ccb','birt','peci','ccw','pum','people tools',
                        'peoplesoft implementation','peoplesoft components',
                        'peoplesoft dba','peoplesoft admin','peoplesoft admin/dba','peopleSoft fscm', 
                        'peopletoolsupgrade','peopletools upgrade','process scheduler servers',
                        'peoplesoft hrms','peopleSoft consultant','peopledoft cloud',
                        'PeopleSoft migrations','eoplesoft Testing Framework','pure internet architecture',
                        'sql','sql server', 'ms sql server','msbi', 'sql developer', 'ssis','ssrs',
                        'ssms','t-sql','tsql','Razorsql', 'razor sql','triggers','powerbi','power bi',
                        'oracle sql', 'pl/sql', 'pl\sql','oracle', 'oracle 11g','oledb','cte','ddl',
                        'dml','etl','mariadb','maria db','reactjs', 'react js', 'react js developer', 'html', 
                        'css3','xml','javascript','html5','boostrap','jquery', 'redux','php', 'node js',
                        'nodejs','apache','netbeans','nestjs','nest js','react developer','react hooks',
                        'jenkins']

# Function to extract text from resume
def getText(filename):
      
    # Create empty string 
    fullText = ''
    if filename.endswith('.docx'):
        doc = docx.Document(filename)
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
            
           
    elif filename.endswith('.pdf'):  
        with open(filename, "rb") as pdf_file:
            pdoc = PyPDF2.PdfFileReader(filename)
            number_of_pages = pdoc.getNumPages()
            page = pdoc.pages[0]
            page_content = page.extractText()
             
        for paragraph in page_content:
            fullText =  fullText + paragraph
            
    else:
        import aspose.words as aw
        output = aw.Document()
        # Remove all content from the destination document before appending.
        output.remove_all_children()
        input = aw.Document(filename)
        # Append the source document to the end of the destination document.
        output.append_document(input, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)
        output.save("Output.docx");
        doc = docx.Document('Output.docx')
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
        fullText = fullText[79:]
         
    return (fullText)

# Function to remove punctuation and tokenize the text
def tokenText(extText):
   
    # Remove punctuation marks
    punc = '''!()-[]{};:'"\,.<>/?@#$%^&*_~'''
    for ele in extText:
        if ele in punc:
            puncText = extText.replace(ele, "")
            
    # Tokenize the text and remove stop words
    stop_words = set(stopwords.words('english'))
    puncText.split()
    word_tokens = word_tokenize(puncText)
    TokenizedText = [w for w in word_tokens if not w.lower() in stop_words]
    TokenizedText = []
  
    for w in word_tokens:
        if w not in stop_words:
            TokenizedText.append(w)
    return(TokenizedText)

# Function to read the tokenized text and search for the key words to dermine the Role Applied for
def roleApplied (Text):
    
    # covert the text to lower case
    for i in range(len(Text)):
        Text[i] = Text[i].lower()
    
    # Obtain the scores for each area
    for area in terms.keys():
        if area == 'WorkDay ERP':
            for word in terms[area]:
                if word in Text:
                    role = area
                    return (role)
                
        elif area == 'Peoplesoft':
            for word in terms[area]:
                if word in Text:
                    role = area
                    return(role)   
                
        elif area == 'Database Developer':
            for word in terms[area]:
                if word in Text:
                    role =  area
                    return(role)
            
        elif area == 'Java Developer':
             for word in terms[area]:
                if word in Text:
                    role = area
                    return(role)
        else:
            role = "Fresher"
            return(role)
        
# Function to extract Name and contact details
def contactDetails(Text):
    name = ''  
    for i in range(0,3):
        name = " ".join([name, Text[i]])
    return(name)

# Function to extract experience details
def expDetails(Text):
    global sent
   
    Text = Text.split()
   
    for i in range(len(Text)-2):
        Text[i].lower()
        
        if Text[i] ==  'years':
            sent =  Text[i-2] + ' ' + Text[i-1] +' ' + Text[i] +' '+ Text[i+1] +' ' + Text[i+2]
            l = re.findall('\d*\.?\d+', sent)
            for i in l:
                a = float(i)
            return (round(a,2))
            #return (sent)
        
        
# Function to extract skill set details
def skillSet(Text):
    t = []
    for i in range(len(Text)):
        if Text[i] in allTerms:
            if Text[i] in t:
                continue
            t.append(Text[i]) 
    return(t)


st.sidebar.header("Resume Classification App")
path = st.sidebar.text_input('Enter the resumes folder path')

ResumeText = pd.DataFrame([], columns=['Name', 'Exp_years', 'SkillSet', 'RoleApplied'])

for filename in os.listdir(path):
    filename = os.path.join(path, filename)
    extText = getText(filename)
    tokText = tokenText(extText)
    role = roleApplied(tokText)
    Name = contactDetails(tokText)
    experience = expDetails(extText)
    skills = skillSet(tokText)
    NewRow = [Name, experience, skills, role]
    ResumeText.loc[len(ResumeText)] = NewRow
    # st.dataframe(ResumeText)
    java = (ResumeText["RoleApplied"] == "Java Developer")
    # javares = ResumeText[java]
    workday = (ResumeText["RoleApplied"] == "WorkDay ERP")
    peosoft = (ResumeText["RoleApplied"] == "Peoplesoft")
    dbms = (ResumeText["RoleApplied"] == "Database Developer")

# this is the main function in which we define our webpage
def main():
    if st.sidebar.button("Display Numbers"):
        st.subheader("No of Resumes Received")
        num = pd.DataFrame(ResumeText['RoleApplied'].value_counts())
        num['Category'] = num.index
        num.set_axis(['No of Resumes', 'Category'], axis='columns', inplace=True)
        num.reset_index(inplace=True, drop=True)
        num = num[['Category', 'No of Resumes']]
        st.dataframe(num)

        base = alt.Chart(num).encode(theta=alt.Theta(
        "No of Resumes:Q", stack=True), color=alt.Color("Category:N", legend=None))
        pie = base.mark_arc(outerRadius=120)
        text = base.mark_text(radius=165, size=15).encode(text="Category:N")
        c = pie + text
        st.altair_chart(c, use_container_width=True)


    if st.sidebar.button("All Resumes"):
        st.subheader("ALL RESUMES")
    exp = st.slider('Select Experience', 0.0, 20.0, 0.1)
    expres = ResumeText[ResumeText['Exp_years'] >= exp]
    st.dataframe(expres)


    if st.sidebar.button("JAVA Resumes"):
        st.subheader("JAVA RESUMES")
    exp1 = st.slider('Select JAVA Experience', 0.0, 20.0, 0.1)
    javadf = ResumeText[java]
    expres1 = javadf[javadf['Exp_years'] >= exp1]
    st.dataframe(expres1)

        #st.dataframe(ResumeText[java])
    if st.sidebar.button("DBMS Resumes"):
        st.subheader("DBMS RESUMES")
    exp2 = st.slider('Select DBMS Experience', 0.0, 20.0, 0.1)
    dbmsdf = ResumeText[dbms]
    expres2 = dbmsdf[dbmsdf['Exp_years'] >= exp2]
    st.dataframe(expres2)

    if st.sidebar.button("Peoplesoft Resumes"):
        st.subheader("Peoplesoft Resumes")
    exp3 = st.slider('Select Peoplesoft Experience', 0.0, 20.0, 0.1)
    psoftdf = ResumeText[peosoft]
    expres3 = psoftdf[psoftdf['Exp_years'] >= exp3]
    st.dataframe(expres3)

    if st.sidebar.button("Workday Resumes"):
        st.subheader("Workday Resumes")
    exp4 = st.slider('Select Workday Experience', 0.0, 20.0, 0.1)
    wdaydf = ResumeText[workday]
    expres4 = wdaydf[wdaydf['Exp_years'] >= exp4]
    st.dataframe(expres4)



if __name__=='__main__': 
    main()

